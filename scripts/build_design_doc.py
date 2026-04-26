"""기술 설계서 (Technical Design Document) 생성 스크립트.

`python scripts/build_design_doc.py` 로 실행하면 프로젝트 루트에
`Starlab_Code_기술설계서.docx` 가 만들어집니다.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


KOREAN_FONT = "맑은 고딕"
MONO_FONT = "Consolas"


def _set_korean_font(run, size: float | None = None, bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = KOREAN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rFonts.set(qn("w:ascii"), KOREAN_FONT)
    rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    sizes = {0: 24, 1: 18, 2: 14, 3: 12}
    _set_korean_font(run, size=sizes.get(level, 12), bold=True, color=(31, 73, 125) if level <= 1 else (0, 0, 0))
    if level == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12 if level >= 2 else 18)
    paragraph.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_korean_font(run, size=11, bold=bold)
    run.italic = italic
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.4


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        _set_korean_font(run, size=11)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.35


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    run = paragraph.add_run(text)
    run.font.name = MONO_FONT
    rPr = run._element.get_or_add_rPr()
    from docx.oxml import OxmlElement

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), MONO_FONT)
    rFonts.set(qn("w:hAnsi"), MONO_FONT)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(40, 40, 40)
    paragraph.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        _set_korean_font(run, size=10.5, bold=True, color=(255, 255, 255))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for cell, text in zip(cells, row):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            _set_korean_font(run, size=10)
    doc.add_paragraph()  # spacer


def build() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(11)

    # Cover
    add_heading(doc, "Starlab Code 기술 설계서", level=0)
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("중·고등학생 알고리즘 학습 / 수업 운영 플랫폼 MVP")
    _set_korean_font(run, size=12, color=(80, 80, 80))
    cover2 = doc.add_paragraph()
    cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover2.add_run("문서 버전 1.0  ·  작성일 2026-04-26")
    _set_korean_font(run, size=10, color=(120, 120, 120))
    doc.add_paragraph()

    # 1. 개요
    add_heading(doc, "1. 시스템 개요", level=1)
    add_paragraph(
        doc,
        "Starlab Code 는 중·고등학생을 대상으로 하는 알고리즘 문제 풀이 및 수업 운영 플랫폼이다. "
        "선생님이 문제를 출제하고 반(class) 단위로 과제를 일괄 배정하면, 학생은 브라우저에서 코드를 작성·제출하고 "
        "테스트케이스 단위 채점 결과를 실시간 진행 바로 확인한다. 운영 측면에서는 학급 단위 완료율, 30일 활동 히트맵, "
        "실시간 제출 피드 등 수업 모니터링 기능을 함께 제공한다.",
    )
    add_heading(doc, "1.1 핵심 사용 시나리오", level=2)
    add_bullets(
        doc,
        [
            "메인 선생님(primary teacher)이 학교/학원 단위로 일반 선생님 계정을 발급한다.",
            "선생님은 학생 계정을 반(class_name)과 함께 만들고, 문제를 출제한 뒤 반 단위로 과제를 배정한다.",
            "학생은 문제은행에서 문제를 풀고, 채점 결과를 테스트케이스별 진행 바로 실시간 확인한다.",
            "선생님은 실시간 제출 피드로 새 제출을 4초 주기로 받아보고, 반별 완료율 드릴다운으로 진척을 점검한다.",
        ],
    )

    add_heading(doc, "1.2 설계 목표", level=2)
    add_bullets(
        doc,
        [
            "교실에서 30~40명이 동시에 제출해도 흔들리지 않을 정도의 단일 노드 처리 성능 확보.",
            "교사가 별도 인프라 지식 없이도 무료 클라우드(Render + Supabase + Cloudflare Pages) 위에 배포 가능.",
            "MVP 단계에서는 신뢰된 환경(학원/교실)을 가정하되, 코드 실행이 호스트를 망가뜨리지 않도록 리소스 가드 적용.",
            "프론트엔드는 Vite 정적 번들 1개로 끝나도록 단순화하여 CDN 한 곳만 있어도 운영 가능하도록 한다.",
            "학생 관점에서 ‘제출 후 결과를 한참 기다리는’ 경험을 없애기 위해 NDJSON 스트리밍 채점을 1급 시민으로 다룬다.",
        ],
    )

    # 2. 전체 아키텍처
    add_heading(doc, "2. 전체 아키텍처", level=1)
    add_paragraph(
        doc,
        "전형적인 3-tier 구조이지만, 채점기는 백엔드 프로세스 내부의 서브프로세스로 동작한다. "
        "외부 큐(Worker)나 Redis 같은 의존성 없이 단일 FastAPI 인스턴스로 학습용 워크로드를 처리하는 것이 의도된 설계이다.",
    )
    add_code(
        doc,
        "[Browser]  ──HTTPS/JWT──>  [FastAPI (Uvicorn)]  ──SQLModel/SQLAlchemy──>  [SQLite or PostgreSQL]\n"
        "                                  │\n"
        "                                  └── asyncio.Semaphore + ThreadPool ──> subprocess(python -I)\n"
        "                                                                          │\n"
        "                                                                          └── stdout(NDJSON) → 학생 화면",
    )

    add_heading(doc, "2.1 컴포넌트 책임 분리", level=2)
    add_table(
        doc,
        ["컴포넌트", "책임"],
        [
            ["Frontend (React + Vite)", "단일 페이지 앱. 로그인/대시보드/문제풀이/계정관리/과제배정/실시간피드 모든 뷰."],
            ["FastAPI 라우터", "REST API + NDJSON 스트리밍 엔드포인트. JWT 인증 의존성으로 권한 분기."],
            ["SQLModel 도메인", "User · Category · Problem · TestCase · Assignment · Submission 6개 핵심 엔티티."],
            ["Judge 서비스", "Python subprocess + POSIX rlimit. 동시성은 BoundedSemaphore 로 제한."],
            ["Seed 모듈", "기동 시 카테고리/문제 100개/테스트케이스 5,000건/데모 계정 incremental 시딩."],
        ],
    )

    add_heading(doc, "2.2 요청 흐름 — 코드 제출", level=2)
    add_bullets(
        doc,
        [
            "① 학생이 ‘제출’을 누르면 fetch 가 POST /problems/{id}/submit/stream 으로 코드와 assignment_id 를 보낸다.",
            "② FastAPI 가 문제·테스트케이스를 조회하고, 채점 generator 를 StreamingResponse 로 감싼다.",
            "③ judge.run_code_iter() 가 테스트마다 subprocess 를 띄우고 결과 dataclass 를 yield 한다.",
            "④ 라우터는 yield 값을 NDJSON {kind:start|result|done} 로 직렬화하여 전송한다.",
            "⑤ 마지막 done 단계에서 별도 세션을 열어 Submission 을 INSERT 하고 submission_id 를 추가로 흘려준다.",
            "⑥ 프론트는 ReadableStream.getReader() 로 한 줄씩 파싱하여 테스트케이스 진행 바를 갱신한다.",
        ],
    )

    # 3. 기술 스택과 선택 이유
    add_heading(doc, "3. 기술 스택과 선택 이유", level=1)

    add_heading(doc, "3.1 한눈에 보기", level=2)
    add_table(
        doc,
        ["계층", "선택", "선택 이유 요약"],
        [
            ["프론트엔드 프레임워크", "React 18 + TypeScript 5.6", "단일 SPA + 타입 안전성. 팀에 가장 익숙한 조합."],
            ["빌드 도구", "Vite 5", "차가운 시작이 빠르고, 정적 dist/ 한 폴더로 끝나 CDN 배포가 단순."],
            ["백엔드 프레임워크", "FastAPI", "타입 힌트 기반 의존성 주입과 자동 OpenAPI, NDJSON StreamingResponse 가 1급 시민."],
            ["서버 런타임", "Uvicorn (ASGI)", "FastAPI 표준 동행. 비동기 + 스레드풀 디스패치로 IO·CPU 혼합에 적합."],
            ["ORM", "SQLModel", "Pydantic 모델과 SQLAlchemy 테이블을 한 클래스에서 정의 → 모델 중복 제거."],
            ["데이터베이스", "SQLite (개발) / PostgreSQL (운영)", "동일 SQLAlchemy URL 만 바꾸면 됨. 무료 운영은 Supabase 의 Postgres."],
            ["인증", "JWT(HS256) + pbkdf2_sha256", "클라이언트가 토큰을 들고 다니는 stateless 구조. 비밀번호 해시는 OS 의존성 없는 pbkdf2."],
            ["채점기", "subprocess + POSIX rlimit", "Docker 등 컨테이너 의존성 없이 단일 호스트에서 즉시 동작. CPU·메모리 강제는 rlimit."],
            ["배포", "Render Free + Supabase + Cloudflare Pages", "월 0원으로 학교/학원 단위 운영이 가능한 조합."],
        ],
    )

    add_heading(doc, "3.2 React + TypeScript 를 선택한 이유", level=2)
    add_bullets(
        doc,
        [
            "교사·학생용 UI 가 모두 한 페이지 앱(SPA)이고, 화면 전환은 view 상태로 분기하면 충분하다.",
            "TypeScript 로 모델 타입을 백엔드 응답과 1:1 로 정의해, 새 필드 누락 같은 실수를 컴파일 단계에서 잡는다.",
            "Next.js 같은 프레임워크가 주는 SSR·라우팅 기능은 본 서비스에 불필요하다. 교실 내 사용자에게는 첫 로드보다 인터랙션 응답성이 더 중요하다.",
            "단일 파일 구조(App.tsx)는 MVP 속도를 위해 의도적으로 유지하되, 향후 기능 분할 시 컴포넌트 추출이 가능하도록 props 기반 설계만 준수했다.",
        ],
    )

    add_heading(doc, "3.3 Vite 를 선택한 이유", level=2)
    add_bullets(
        doc,
        [
            "esbuild 기반 dev server 로 첫 기동이 1초 미만 → 수업 중 빠른 수정·배포에 유리.",
            "결과물이 dist/ 정적 자산뿐이라 Cloudflare Pages 의 무료 CDN 에 그대로 올릴 수 있다.",
            "환경변수 VITE_API_BASE_URL 만으로 백엔드 주소를 빌드타임 고정 → CORS·도메인 분리 운영이 깔끔하다.",
            "_redirects 한 줄로 Pages 의 SPA fallback 이 해결되므로 별도 호스팅 설정이 필요 없다.",
        ],
    )

    add_heading(doc, "3.4 FastAPI 를 선택한 이유", level=2)
    add_bullets(
        doc,
        [
            "Python 으로 채점기를 작성해야 하므로 서버도 Python 으로 통일해 인프라/언어 매트릭스를 줄였다.",
            "타입 힌트 기반 의존성 주입(Depends)으로 ‘로그인 사용자’, ‘선생님 권한’ 같은 가드를 함수 시그니처로 표현 가능.",
            "StreamingResponse 가 표준이라 NDJSON 채점 스트림을 별도 라이브러리 없이 구현했다.",
            "자동 OpenAPI(/docs)로 프론트 개발자와 별도 문서 없이도 API 계약을 공유할 수 있다.",
            "Flask 대비 비동기 / 검증 / 직렬화가 기본 제공되어 부수적인 코드를 줄일 수 있다.",
        ],
    )

    add_heading(doc, "3.5 SQLModel 을 선택한 이유", level=2)
    add_bullets(
        doc,
        [
            "Pydantic 검증 모델과 SQLAlchemy 매핑을 동일 클래스에서 정의하여 ‘DB 모델 ↔ API 스키마’ 중복을 제거한다.",
            "FastAPI 와 같은 작성자(Sebastián Ramírez)가 만든 라이브러리라 결합 시 타입·문서가 일관된다.",
            "SQLAlchemy 코어를 그대로 쓸 수 있어, select().where().order_by() 형태로 익숙한 쿼리를 유지한다.",
            "마이그레이션은 MVP 단계에서 db.py 의 ALTER TABLE 자동 보정으로 충분하며, Alembic 도입은 운영 단계로 미룬다.",
        ],
    )

    add_heading(doc, "3.6 SQLite + PostgreSQL 듀얼 백엔드", level=2)
    add_bullets(
        doc,
        [
            "로컬 개발은 파일 1개(starlab_code_mvp.db)로 끝나는 SQLite 가 압도적으로 편하다 — 설치, 백업, 초기화가 즉시.",
            "운영은 Supabase Free Postgres 로 옮겨서 동시 쓰기·풀링·백업 안정성을 확보한다.",
            "config.py 의 _normalize_database_url() 이 postgres:// → postgresql+psycopg:// 자동 변환으로 Supabase 의 ‘Session pooler’ 문자열을 그대로 받는다.",
            "SQLite 모드에서는 db.py 의 PRAGMA(WAL, synchronous=NORMAL, busy_timeout=5000ms) 설정으로 30~40명 동시 풀이에서 락 충돌을 줄였다.",
        ],
    )

    add_heading(doc, "3.7 인증 — JWT + pbkdf2_sha256", level=2)
    add_bullets(
        doc,
        [
            "교실 환경에서 SSO·OAuth 를 강제하기 어렵기 때문에 username + password 단순 로그인 + JWT 로 시작했다.",
            "토큰 만료(STARLAB_TOKEN_MINUTES, 기본 24시간)는 환경변수로 조정 가능하다. HS256 비밀키는 Render render.yaml 의 generateValue 로 자동 생성.",
            "비밀번호 해시는 bcrypt 가 아닌 pbkdf2_sha256 을 사용한다 — Render Free 의 musl/glibc 환경에서 bcrypt 의 C 모듈 빌드가 실패하는 사례를 피하기 위해서다.",
            "역할(UserRole)과 메인선생님 여부(is_primary_teacher)는 토큰이 아니라 DB 사용자 레코드에서 매 요청마다 조회한다 — 권한 변경이 즉시 반영되도록.",
        ],
    )

    add_heading(doc, "3.8 채점기 — subprocess + rlimit", level=2)
    add_bullets(
        doc,
        [
            "‘공개 인터넷에서 임의 코드 실행’이 아닌 ‘교사가 만든 반 안에서만 코드 실행’이 운영 모델이다. 따라서 Docker / gVisor 격리 없이 subprocess 로 시작했다.",
            "Linux/macOS 에서는 RLIMIT_CPU(4초), RLIMIT_AS(256MB), RLIMIT_FSIZE, RLIMIT_NPROC, RLIMIT_CORE 를 prefork 단계에서 설정한다.",
            "stdin 주입과 출력 캡처는 별도 스레드 + 4KB 청크 + 256KB 캡으로 처리해, 무한 출력 코드가 메모리를 폭주시키는 것을 막는다.",
            "동시성은 asyncio 가 아니라 threading.BoundedSemaphore(STARLAB_JUDGE_CONCURRENCY) 로 강제한다 — 동기 subprocess.wait() 와 호환되는 가장 단순한 직렬화 방법이다.",
            "Windows 는 resource 모듈 부재로 rlimit 이 적용되지 않으므로 ‘로컬 개발 전용’으로 명시했다.",
        ],
    )

    add_heading(doc, "3.9 NDJSON 스트리밍을 채택한 이유", level=2)
    add_bullets(
        doc,
        [
            "테스트케이스가 50개라면 일괄 결과를 기다리는 학생 입장에서 ‘제출 후 정지’ 경험이 발생한다.",
            "WebSocket 은 Render Free / Cloudflare 조합에서 추가 설정·연결 유지 비용이 있어 피했다.",
            "Server-Sent Events 도 후보였으나, 단방향 페이로드만 필요하고 OpenAPI 통합이 떨어진다.",
            "결과적으로 ‘fetch + ReadableStream + JSON 라인’ 조합이 가장 단순하다 — 백엔드는 yield, 프론트는 split('\\n').",
        ],
    )

    add_heading(doc, "3.10 배포 스택 — Render + Supabase + Cloudflare Pages", level=2)
    add_bullets(
        doc,
        [
            "Render Free 는 render.yaml 한 파일만 있으면 GitHub 연동 후 자동 배포된다. 15분 idle sleep 단점은 학교 수업 시작 직전 한 번만 깨우면 되는 시나리오와 잘 맞는다.",
            "Supabase Free Postgres 는 IPv4 에서 동작하는 Session pooler 문자열을 제공해 Render 와 자연스럽게 결합된다.",
            "Cloudflare Pages 는 정적 dist/ 만 받고 글로벌 CDN 으로 서빙하므로, React 번들의 첫 로드 지연을 최소화한다.",
            "전 구간 합쳐 월 0원 운영이 가능 — 학교/학원 MVP 도입 비용을 사실상 ‘없음’으로 만들었다.",
        ],
    )

    # 4. 데이터 모델
    add_heading(doc, "4. 데이터 모델", level=1)
    add_paragraph(
        doc,
        "도메인 엔티티는 6개로 단순하게 설계했다. 모든 모델은 backend/app/models.py 에서 SQLModel 로 정의되며, "
        "동일 모듈에서 응답 전용 Pydantic 클래스(`*Read`, `*Card`, `*Detail`)도 함께 선언한다.",
    )

    add_table(
        doc,
        ["엔티티", "주요 필드", "관계 / 역할"],
        [
            ["User", "id, username, hashed_password, role, class_name, primary_teacher_id, created_by_teacher_id, is_primary_teacher", "교사·학생 통합. 메인 선생님 → 일반 선생님 → 학생의 3계층 위계 표현."],
            ["Category", "id, name, description", "10개의 알고리즘 분류 시드(예: 정렬·탐색, 그래프 기초)."],
            ["Problem", "id, title, statement, category_id, difficulty, time_limit_seconds, memory_limit_mb, starter_code_python", "Category(N:1). 문제별 제한과 스타터 코드 포함."],
            ["TestCase", "id, problem_id, input_data, expected_output, is_public", "Problem(N:1). 문제당 10~50개 강제. is_public 으로 공개/비공개 구분."],
            ["Assignment", "id, title, problem_id, teacher_id, student_id, assignment_type, classroom_label", "‘1행 = 1학생’. 반 단위 배정 시 학생 수만큼 행이 생성된다."],
            ["Submission", "id, problem_id, user_id, assignment_id, code, status, passed_tests, total_tests, runtime_ms", "제출 이력. (assignment_id, user_id) 쌍으로 진척도 집계."],
        ],
    )

    add_heading(doc, "4.1 선생님 위계 표현", level=2)
    add_bullets(
        doc,
        [
            "is_primary_teacher = true → 메인 선생님: 다른 선생님 계정을 발급할 수 있는 관리자 역할.",
            "primary_teacher_id → 같은 조직(학원/학교) 안에서 ‘소속 메인 선생님’ 을 가리킨다.",
            "created_by_teacher_id → 학생을 직접 생성한 선생님. 학생의 데이터 가시성은 이 필드 기준으로 제한된다.",
            "이 3개 컬럼은 db.py 의 _run_schema_migrations() 가 기동 시 ALTER TABLE 로 자동 추가한다.",
        ],
    )

    add_heading(doc, "4.2 인덱스/성능 설계", level=2)
    add_bullets(
        doc,
        [
            "User: username UNIQUE 인덱스 + class_name / primary_teacher_id / created_by_teacher_id 보조 인덱스.",
            "Category.name UNIQUE 인덱스로 시드 incremental 삽입 시 빠른 존재 검사.",
            "Submission 은 assignment_id · user_id 기반 조회가 잦아, 가능한 in_() 한 번으로 묶어 N+1 을 회피하도록 라우터에서 설계했다.",
            "Problem 검색은 LIKE %title% 정도로 충분하다 — 전체 문제 수가 수백 단위라 풀스캔 비용이 무시 가능.",
        ],
    )

    # 5. API 설계
    add_heading(doc, "5. API 설계", level=1)
    add_paragraph(
        doc,
        "모든 라우트는 backend/app/main.py 에 정의되어 있으며, OpenAPI 문서는 /docs 에서 자동 생성된다. "
        "권한은 FastAPI Depends 의 get_current_user / require_teacher 두 의존성으로 단순하게 분기한다.",
    )

    add_table(
        doc,
        ["분류", "대표 라우트", "설명"],
        [
            ["인증", "POST /auth/token, GET /auth/me", "username + password 로 JWT 발급, 현재 사용자 조회."],
            ["계정 관리", "POST /users/teachers, POST /users/students", "메인 선생님은 선생님 발급, 선생님은 학생 발급."],
            ["문제", "GET/POST/PUT /problems[/{id}]", "교사 전용 작성·수정. 학생/공개에는 공개 테스트만 노출."],
            ["과제", "POST /assignments, GET /assignments/groups[/detail]", "반 단위 일괄 배정 + 그룹별 완료율/드릴다운."],
            ["채점", "POST /problems/{id}/run, /submit (+ /stream)", "공개 테스트 / 전체 테스트 채점. 스트리밍 변형은 NDJSON."],
            ["모니터링", "GET /submissions/feed?since_id=", "선생님 전용 4초 폴링 피드. 신규 행만 반환."],
            ["대시보드", "GET /dashboard, /classrooms, /students, /teachers", "역할별 요약 지표 + 반 목록 + 명부."],
        ],
    )

    add_heading(doc, "5.1 NDJSON 응답 스키마", level=2)
    add_code(
        doc,
        '{"kind":"start","total":50}\n'
        '{"kind":"result","index":0,"status":"passed","runtime_ms":18, ...}\n'
        '{"kind":"result","index":1,"status":"wrong_answer", ...}\n'
        '{"kind":"done","status":"wrong_answer","passed_tests":49,"total_tests":50,"runtime_ms":124,"submission_id":1234}',
    )

    # 6. 보안 / 운영 고려
    add_heading(doc, "6. 보안 및 운영 고려", level=1)
    add_table(
        doc,
        ["항목", "현재 대응", "운영 단계 권장"],
        [
            ["인증", "JWT HS256 + pbkdf2_sha256", "비밀키 외부화(이미 적용). 토큰 만료 단축, refresh 분리 도입."],
            ["권한", "역할 + primary_teacher_id 격리", "감사 로그(audit log) 추가."],
            ["코드 실행", "subprocess + rlimit + 출력 캡", "Docker / gVisor / firejail 등 컨테이너 격리. 네트워크 차단."],
            ["입력 검증", "Pydantic + 테스트케이스 10~50개 강제", "파일 업로드 도입 시 MIME/크기 화이트리스트."],
            ["CORS", "STARLAB_ALLOW_ORIGINS 화이트리스트", "운영 도메인만 허용, 와일드카드 금지."],
            ["DB", "SQLite WAL / Postgres 풀 5+5", "PgBouncer/Session Pooler 활용. 백업 자동화."],
            ["DoS 보호", "(없음)", "slowapi 등 레이트 리밋. /submit/stream 은 사용자별 큐 도입."],
        ],
    )

    # 7. 성능 / 확장성
    add_heading(doc, "7. 성능과 확장성", level=1)
    add_bullets(
        doc,
        [
            "FastAPI 의 default thread limiter 를 STARLAB_THREADPOOL_SIZE(기본 64)로 끌어올려, 동기 ORM 호출이 많은 라우트에서 워커 부족이 일어나지 않도록 했다.",
            "채점 동시성은 STARLAB_JUDGE_CONCURRENCY(기본 4) 세마포어로 직렬화한다 — Render Free 의 0.5 vCPU 기준 안전한 값이다.",
            "프론트 실시간 피드는 4초 폴링 + since_id 증분 조회로, Render Free 의 1Mbps 대역폭을 거의 점유하지 않는다.",
            "Submission 조회는 in_() 일괄 + dict 매핑 패턴으로 학생 30명 기준 N+1 을 1회 쿼리로 압축했다.",
            "수직 확장 한계는 채점 큐 — 동시 100명 이상 제출이 예상되면 별도 워커(Celery/RQ) 분리가 다음 단계.",
        ],
    )

    # 8. 프론트엔드 설계
    add_heading(doc, "8. 프론트엔드 설계", level=1)
    add_paragraph(
        doc,
        "프론트엔드는 의도적으로 단일 컴포넌트 트리(App.tsx, 약 4,400줄)로 작성했다. "
        "MVP 단계에서 라우터/상태관리 라이브러리 도입의 부수적인 복잡성보다, view 상태값 1개로 분기하는 단순함의 가치가 크다고 판단했다.",
    )
    add_heading(doc, "8.1 화면 분기", level=2)
    add_table(
        doc,
        ["view 키", "화면", "사용자"],
        [
            ["home", "역할별 대시보드(학생: 활동 히트맵 / 선생님: 반 활동 요약)", "공통"],
            ["problems", "문제은행 (분류·난이도 필터, solved.ac 풍 티어 색)", "공통"],
            ["solve", "문제 풀이 + NDJSON 스트리밍 채점 UI", "공통"],
            ["submissions", "내 제출 이력", "학생"],
            ["live", "실시간 제출 피드 (4초 폴링 + 신규 행 하이라이트)", "선생님"],
            ["assignments", "과제 배정/현황 + 반별 드릴다운", "선생님"],
            ["accounts", "계정 관리 (선생님 / 학생 서브탭)", "선생님"],
            ["manage", "문제 생성·수정 (테스트케이스 10~50개)", "선생님"],
        ],
    )
    add_heading(doc, "8.2 스트리밍 클라이언트", level=2)
    add_paragraph(
        doc,
        "executeStream() 헬퍼가 fetch + ReadableStream.getReader() 로 청크를 받아 \\n 단위로 분할한다. "
        "각 라인은 JSON.parse 후 kind 에 따라 진행 바 / 결과 / 완료 상태를 갱신한다. "
        "에러가 나도 done 메시지가 오지 않은 경우 ‘네트워크 끊김’ 으로 표시되도록 한 번의 try/finally 로 감쌌다.",
    )

    # 9. 시드 / 데모 데이터
    add_heading(doc, "9. 시드 데이터 전략", level=1)
    add_bullets(
        doc,
        [
            "기동 시 on_startup → seed_initial_data() 가 동기 실행된다. 운영용 DB 가 비어 있어도 즉시 사용 가능한 상태로 만든다.",
            "카테고리(10) → 문제(100) → 테스트케이스(문제당 50) 순으로 들어가며 총 5,000건의 케이스가 결정론적으로 생성된다.",
            "문제 시드는 ‘제목 기준 incremental 삽입’ 이라, 운영 중 새 문제만 추가되며 기존 사용자 제출 이력은 유지된다.",
            "데모 계정(메인 선생님 1, 일반 선생님 1, 학생 2)도 함께 ensure 되어 첫 데모 세션을 즉시 시연할 수 있다.",
        ],
    )

    # 10. 향후 작업 / 한계
    add_heading(doc, "10. 알려진 한계와 향후 작업", level=1)
    add_table(
        doc,
        ["영역", "현재 한계", "다음 단계"],
        [
            ["채점 격리", "subprocess + rlimit 수준. 임의 코드의 네트워크 차단 안 됨.", "Docker --network=none / gVisor / firejail 래퍼. 또는 Oracle Cloud Always Free + 컨테이너."],
            ["언어 지원", "Python 만 채점 가능 (language 필드만 분리됨).", "C/C++ 컴파일 단계 추가, 컨테이너 이미지에 gcc 사전 설치."],
            ["채점 큐", "in-process 세마포어. 같은 인스턴스가 다운되면 진행 중 채점 손실.", "Celery / RQ 분리, Redis 큐, 결과 영속화."],
            ["DB 마이그레이션", "ALTER TABLE 자동 보정만 지원.", "Alembic 도입, 마이그레이션 히스토리 관리."],
            ["레이트 리밋", "없음.", "slowapi 등 도입. /submit 은 사용자별 동시 1건 제한."],
            ["프론트엔드 구조", "App.tsx 단일 파일.", "기능별 컴포넌트/훅 분리, react-router 도입."],
            ["관측성", "stdout 로그만.", "Render 로그 → Logflare / Grafana Loki 연동, /metrics 노출."],
        ],
    )

    # 11. 부록
    add_heading(doc, "11. 부록 — 환경변수 요약", level=1)
    add_table(
        doc,
        ["변수", "기본값", "설명"],
        [
            ["STARLAB_SECRET_KEY", "(개발용 하드코드)", "JWT HS256 서명 키. 운영 전 반드시 외부 주입."],
            ["STARLAB_DATABASE_URL", "sqlite:///starlab_code_mvp.db", "postgres:// 는 자동으로 postgresql+psycopg:// 로 정규화."],
            ["STARLAB_ALLOW_ORIGINS", "localhost:5173/4173", "콤마 구분 CORS 화이트리스트."],
            ["STARLAB_TOKEN_MINUTES", "1440", "JWT 만료 분."],
            ["STARLAB_PRIMARY_TEACHER_*", "main_teacher / 메인 선생님 / ChangeMe1234!", "기동 시 자동 ensure 되는 메인 선생님 계정."],
            ["STARLAB_JUDGE_CONCURRENCY", "4", "동시 채점 서브프로세스 상한."],
            ["STARLAB_JUDGE_CPU_SECONDS", "4", "POSIX rlimit CPU 초."],
            ["STARLAB_JUDGE_MEMORY_BYTES", "268435456", "POSIX rlimit AS(메모리) 바이트."],
            ["STARLAB_MAX_CODE_BYTES", "65536", "제출 코드 크기 상한(64KB)."],
            ["STARLAB_MAX_INPUT_BYTES", "262144", "테스트 입력 크기 상한(256KB)."],
            ["STARLAB_MAX_OUTPUT_BYTES", "262144", "프로세스 출력 크기 상한(256KB)."],
            ["STARLAB_DB_POOL_SIZE / MAX_OVERFLOW / RECYCLE_SECONDS", "5 / 5 / 300", "Postgres 풀 설정 (Render Free + Supabase 권장값)."],
            ["STARLAB_THREADPOOL_SIZE", "64", "FastAPI default thread limiter 토큰 수."],
            ["STARLAB_SEED_DEMO_DATA", "true", "기동 시 데모 시드 실행 여부."],
        ],
    )

    add_heading(doc, "12. 결론", level=1)
    add_paragraph(
        doc,
        "Starlab Code 는 ‘교실 단위 학습용 코드 채점 플랫폼’ 이라는 좁은 도메인에 맞춰, "
        "운영 비용은 0원, 인프라 의존성은 최소, 학습자 경험(스트리밍 채점)은 1급 시민이라는 세 가지 원칙으로 설계되었다. "
        "선택된 기술 — FastAPI · SQLModel · React + Vite · subprocess 채점 — 는 모두 ‘한 사람이 한 호스트에서 운영 가능’ 이라는 MVP 요구를 충족하기 위한 결정이다. "
        "동시에 PostgreSQL 듀얼 백엔드, NDJSON 스트리밍, 메인 선생님 위계처럼, 학교/학원 단위로 확장될 때 필요한 ‘다음 한 걸음’을 위한 여지를 모델 수준에서 미리 마련해 두었다.",
    )

    output_path = Path(__file__).resolve().parent.parent / "Starlab_Code_기술설계서.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    saved = build()
    print(f"saved: {saved}")
